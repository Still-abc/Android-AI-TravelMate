from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = Path(r"E:\Android_work\Wetravel\基于Android的AI智能旅行规划系统（WeTravel）的设计与实现-课程设计报告.docx")


def set_run_font(run, size=None, bold=None, color=None, font="宋体"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_font(paragraph, size=11, font="宋体"):
    for run in paragraph.runs:
        set_run_font(run, size=size, font=font)


def set_cell_text(cell, text, bold=False, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, font="宋体")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.first_child_found_in("w:tcW")
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")


def add_title(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    color = "1F4D78" if level <= 2 else "000000"
    size = {1: 16, 2: 14, 3: 12}.get(level, 11)
    set_run_font(run, size=size, bold=True, color=color, font="黑体")
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(doc, text, first_line=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size=11, font="宋体")
    return p


def add_center(doc, text, size=11, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, font="宋体")
    return p


def add_figure_placeholder(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("【" + caption + "】")
    set_run_font(run, size=10.5, bold=True, color="555555", font="宋体")
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        shade_cell(hdr[i], "E8EEF5")
        if widths:
            set_cell_width(hdr[i], widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                set_cell_width(cells[i], widths[i])
    doc.add_paragraph()
    return table


def add_cover(doc):
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run("《Android开发技术》课程项目设计报告")
    set_run_font(r, size=22, bold=True, font="黑体")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(44)
    r = p.add_run("基于Android的AI智能旅行规划系统（WeTravel）的设计与实现")
    set_run_font(r, size=18, bold=True, font="黑体")
    items = [
        ("学院", "计算机学院"),
        ("专业", ""),
        ("班级", ""),
        ("学号", ""),
        ("姓名", ""),
        ("指导教师", ""),
        ("完成日期", "2026年7月"),
    ]
    table = doc.add_table(rows=len(items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row, (k, v) in zip(table.rows, items):
        set_cell_text(row.cells[0], k, bold=True, size=12)
        set_cell_text(row.cells[1], v, size=12)
        set_cell_width(row.cells[0], 4)
        set_cell_width(row.cells[1], 8)
    doc.add_page_break()


def build_document():
    doc = Document()
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    add_cover(doc)

    add_title(doc, "摘要", 1)
    add_para(doc, "WeTravel是一款面向个人出行场景的Android智能旅行规划系统。项目源码显示，系统以Kotlin为主要开发语言，采用Jetpack Compose构建界面，通过Navigation Compose组织启动、登录、注册、首页、AI规划、行程、天气、地图、收藏、历史、个人中心和设置等页面；业务状态由ViewModel和StateFlow承载，Repository负责封装本地数据库、会话状态、第三方服务和AI接口。项目通过Hilt完成依赖注入，Room承担用户资料、收藏、历史记录和天气、POI、图片缓存，Retrofit与OkHttp负责AI、和风天气、高德地图和Pexels图片服务通信，kotlinx.serialization完成JSON解析。AI能力采用OpenAI兼容的chat/completions接口，PromptManager明确约束旅行计划、天气建议、预算分析和攻略总结的JSON返回结构，AIRepository还实现了字段归一化，保证页面能够稳定展示结构化行程。系统围绕旅行前的信息收集、路线生成、地点检索、天气提醒、图片呈现和个人数据沉淀形成闭环，具有较完整的移动端课程项目特征。本文在分析源码目录、核心类、数据库表、网络请求、AI流程和页面交互的基础上，对系统设计与实现进行说明，并结合构建测试结果总结项目亮点和改进方向。")
    add_para(doc, "关键词：Android；Jetpack Compose；MVVM；Room；AI旅行规划")

    add_title(doc, "第一章 前言", 1)
    add_title(doc, "1.1 项目目的", 2)
    add_para(doc, "本项目的开发目的在于将移动端界面开发、本地数据持久化、网络服务调用和大模型结构化生成能力整合到同一旅行应用中。课程项目并未停留在单一页面展示，而是围绕用户从注册登录、完善资料、搜索城市、查看天气和地点、生成AI行程、查看历史与收藏的连续使用过程展开。源码中的MainActivity以Compose作为入口，AITravelMateApplication启用Hilt，AppNavHost集中管理页面跳转，Repository和ViewModel共同承担业务状态转换，这些实现使项目能够体现Android开发技术课程对架构、页面、数据和网络的综合要求。")
    add_para(doc, "从软件工程角度看，本项目的目标是建立一个职责划分清晰、数据流向明确、可继续扩展的旅行规划系统。用户输入出发地、目的地、日期、预算、人数和兴趣后，PlannerScreen通过AIViewModel调用AIRepository生成结构化旅行计划，GeneratedPlanRepository将结果写入历史记录并供ItineraryScreen展示。首页和地图模块还会并行调用天气、POI和图片接口，使AI规划结果不只是文字回答，而能与真实城市数据、缓存数据和个人收藏产生联动。")

    add_title(doc, "1.2 项目概述", 2)
    add_para(doc, "WeTravel项目位于单一Android应用模块app中，包名为com.example.ai。源码目录下存在ai、thirdparty、database、repository、navigation、ui、component、theme、config等包。ai包负责大模型请求、Prompt管理、AI数据模型和AI状态管理；thirdparty包负责天气、地图、图片、定位等第三方能力；database包定义Room数据库、DAO、缓存实体和迁移脚本；repository包实现用户、会话、收藏、历史等本地业务；ui包按功能页面拆分；component包沉淀可复用Compose组件；theme包统一颜色、字体、形状和间距。")
    add_para(doc, "项目构建文件app/build.gradle.kts显示，系统启用了Android Gradle Plugin、Kotlin Android、Kotlin Compose、Kotlin Serialization、Kapt和Hilt插件，compileSdk与targetSdk均为36，minSdk为24，Java与Kotlin目标版本为17。依赖中包含Compose BOM、Material3、Navigation Compose、Lifecycle、Room、Retrofit、OkHttp、Coil、Coroutines、Hilt等库。该依赖组合与源码中的实现一致，说明项目主要采用现代Android原生技术栈完成，不依赖传统XML布局。")

    add_title(doc, "第二章 需求分析", 1)
    add_title(doc, "2.1 项目背景", 2)
    add_para(doc, "移动旅行应用的关键价值在于减少用户在出发前反复切换多个工具的成本。传统流程中，用户需要分别查询天气、搜索景点、比对酒店、美食、预算和路线，再将结果人工整理成行程表。WeTravel源码围绕这一痛点组织功能：HomeScreen自动定位或搜索城市，并同步展示天气、景点、酒店、美食和图片；PlannerScreen收集用户的预算、日期和兴趣，直接将约束交给AI接口生成行程；ItineraryScreen把结果转换为按天分组的时间线；HistoryScreen保存并重开完整计划。该项目的背景不是单纯展示旅行信息，而是通过本地缓存、第三方数据和AI生成共同完成旅行计划闭环。")

    add_title(doc, "2.2 功能需求", 2)
    modules = [
        ("2.2.1 登录模块", "登录功能由LoginScreen.kt中的LoginViewModel、LoginRepository和SessionRepository共同完成。界面层提供手机号、密码、记住登录等输入控件，并在提交前校验大陆手机号正则和密码长度。LoginRepository先通过UserDao.getByPhoneAndPassword查询本地users表，兼容旧用户password为空的情况，并在登录成功后调用SessionRepository.setLoggedIn写入SharedPreferences中的current_phone和remember_login。启动阶段SplashRepository会读取rememberedPhone并确认用户仍存在，再决定进入首页或登录页。该模块体现了本项目的本地账号体系：账号认证没有接入远程服务器，用户凭据和会话状态均保存在设备端，界面状态以MutableStateFlow维护，Compose通过collectAsStateWithLifecycle自动刷新错误提示、加载状态和跳转结果。"),
        ("2.2.2 注册模块", "注册功能同样位于LoginScreen.kt。RegisterViewModel负责发送验证码、倒计时、验证码校验、密码确认和协议勾选判断。验证码由本地随机生成六位数字，并在页面消息中显示为开发验证码，倒计时通过viewModelScope每秒更新StateFlow。注册通过校验后并未立即写入数据库，而是导航到CompleteProfile路由，将手机号和密码作为路由参数传给ProfileSetupScreen。ProfileViewModel在完善资料页面收集昵称、头像、性别、生日、签名、所在城市和旅行偏好，再调用UserRepository.saveUser写入Room的users表。这样的流程把账号创建和资料完善拆分为两个阶段，保证users表中的核心展示字段在注册完成时一次性形成。"),
        ("2.2.3 首页模块", "首页由HomeScreen和HomeDataViewModel实现。进入首页后，页面先检查定位权限，已授权时调用LocationRepository.detectCurrentCity获取当前城市，未授权时提示用户搜索城市。城市加载完成后，HomeDataViewModel通过coroutineScope并行请求WeatherRepository、MapRepository和ImageRepository，分别获取天气、景点、酒店、美食和城市图片，并把结果合并到HomeDataUiState中。页面展示用户头像昵称、城市搜索栏、城市横幅、收藏城市按钮、快捷AI规划入口、天气摘要、热门景点、酒店推荐和美食推荐。首页还接入AIViewModel，用户点击推荐攻略卡片后会直接生成短途旅行计划并跳转到行程页。该模块承担系统的信息聚合入口，是第三方API、AI生成和个人收藏共同进入用户界面的主要位置。"),
        ("2.2.4 AI旅行规划模块", "AI旅行规划由PlannerScreen、AIViewModel、AIRepository、PromptManager和GeneratedPlanRepository组成。PlannerScreen采集目的地、出发日期、返程日期、总预算、交通预算、住宿预算、餐饮预算、人数和兴趣标签，并在点击生成按钮后调用AIViewModel.generateTravelPlan。AIRepository使用PromptManager.travelPlanPrompt生成system和user消息，要求大模型只返回合法JSON对象，并明确字段必须匹配TravelPlan、TravelDay、Schedule、Budget、Hotel、Food和Tips数据类。请求成功后，AIRepository用kotlinx.serialization解析结果，对字段缺失或命名差异进行normalizeTravelPlanJson处理，再把结构化TravelPlan交给GeneratedPlanRepository保存到历史记录。页面层在aiState.travelPlan更新后触发行程增强和导航，形成从输入约束到结构化展示的完整链路。"),
        ("2.2.5 城市搜索模块", "城市搜索在首页和地图页同时存在。HomeScreen中的搜索栏把用户输入交给HomeDataViewModel.loadCity，MapScreen也可通过initialCity参数加载指定城市。ViewModel对空字符串进行拦截，随后调用loadCityData统一请求城市相关数据。城市不是静态枚举，而是作为天气、POI、图片和AI规划的共同参数在各Repository之间传递。MapRepository.searchPoi把城市名发送到高德place/text接口，WeatherRepository通过高德地理编码或和风天气接口解析位置，ImageRepository把城市与travel city等关键词发送到Pexels。城市搜索的结果不单独存入cities表，而是被拆解为天气缓存、POI缓存和图片缓存，并在首页、地图、详情页和收藏模块复用。"),
        ("2.2.6 天气模块", "天气模块由WeatherScreen、WeatherViewModel和thirdparty.repository.WeatherRepository实现。页面从Navigation参数读取city，默认城市为上海，随后调用WeatherViewModel.loadWeather。Repository先检查cached_weather表中以城市和日期拼接形成的cacheKey，缓存未过期并且空气质量字段有效时直接返回。缓存失效后先调用和风天气实时天气、三日预报和空气质量接口，接口异常时回退到高德天气weatherInfo。返回数据被映射为WeatherForecast，包含城市、日期、天气、温度、最高最低温、湿度、风力、空气质量、降水和图标。WeatherViewModel在天气成功后继续调用AIRepository.generateWeatherAdvice，页面把穿搭、雨伞、防晒、户外建议合并到WeatherCard中展示。"),
        ("2.2.7 地图模块", "地图模块由MapScreen、MapViewModel、LocationRepository和MapRepository组成。MapScreen进入时读取路由中的城市参数，参数为空时尝试请求定位权限并读取系统LocationManager中的最近位置或单次定位结果。LocationRepository将经纬度交给MapRepository.reverseGeocode，使用高德逆地理编码识别城市。MapViewModel.loadCityData通过async并行调用searchScenic、searchHotels和searchFood，分别搜索景点、酒店和美食。页面按照三类POI分区展示名称、类型、地址、评分和电话，并提供打开地图路线按钮。该按钮不是内嵌地图控件，而是构造androidamap://viewMap深链，设备未安装高德地图时回退到https://uri.amap.com链接。"),
        ("2.2.8 景点模块", "景点数据来自MapRepository.searchScenic，其内部调用searchPoi(city,  景点, 110000)，对应高德POI的风景名胜类目。高德返回的AMapPoi会通过toPoi映射为项目内部Poi模型，保留id、name、type、address、tel、location、rating和photoUrl。首页的PoiRow把景点以横向ScenicCard展示，点击后通过AppRoute.Scenic.createRoute把名称、地址、电话、经纬度、评分、图片、类型和分类作为路由参数传给ScenicDetailScreen。详情页优先使用高德POI图片，其次使用静态地图图片，再使用预置Unsplash旅行图片。景点模块没有单独写入scenic_spots表，实际运行以高德POI和cached_poi缓存为主，历史Room实体中的Scenic更像预留的结构化数据表。"),
        ("2.2.9 酒店模块", "酒店模块同样基于高德POI查询实现。MapRepository.searchHotels调用searchPoi(city, 酒店, 100000)，返回住宿类地点。首页展示酒店推荐列表，地图页展示附近酒店分区，详情页通过ScenicDetailScreen以category参数区分Hotel类别，并使用HotelFallbackImages提供兜底图片。AI旅行规划中还有另一类酒店数据，即AIModels.kt中的Hotel数据类，字段包括name、location、pricePerNight和reason。ItineraryScreen在旅行提示中展示AI推荐的住宿理由和预算概览。由此可见，项目中酒店模块存在两条来源：城市浏览场景依赖高德POI，AI行程场景依赖大模型生成的结构化hotels数组，两者在UI上共同服务住宿推荐。"),
        ("2.2.10 收藏模块", "收藏模块围绕城市收藏实现。HomeScreen在城市加载完成后根据HomeDataUiState.favoriteCities判断当前城市是否已收藏，点击收藏按钮会调用HomeDataViewModel.toggleCurrentCityFavorite，再进入FavoriteRepository.toggleCityFavorite。该方法读取SessionRepository.currentPhone，要求用户处于登录状态，并以userPhone、targetType=city和targetId=城市名查询favorites表。未收藏时写入Favorite实体，记录标题、图片、创建时间和用户手机号；已收藏时删除对应记录。FavoriteScreen通过FavoriteViewModel订阅FavoriteRepository.observeCityFavorites，使用Flow随数据库变化自动刷新列表，点击收藏城市后跳转到对应城市的地图页面。"),
        ("2.2.11 历史记录模块", "历史记录模块用于保存AI生成的完整旅行计划。AI规划成功后，AIViewModel调用GeneratedPlanRepository.updatePlan，该方法一方面把最新TravelPlan放入MutableStateFlow供ItineraryScreen实时展示，另一方面调用HistoryRepository.addPlanHistory写入history表。History实体记录目标城市、标题、描述、天数、行程数量、酒店数量、美食数量以及planJson。HistoryScreen通过GeneratedPlanRepository.observeHistory把History映射为GeneratedPlanHistory，并以时间线形式展示最近生成记录。点击历史记录会调用openHistoryPlan，从planJson反序列化为TravelPlan并重新写入latestPlan，再跳转到ItineraryScreen；长按记录触发AlertDialog确认删除。该模块使AI结果具有可恢复性，而不是一次性页面状态。"),
        ("2.2.12 设置模块", "设置模块由SettingsScreen和SettingsViewModel实现，当前页面提供Dark Mode、Notifications、Cache、Language、About和Feedback六类设置行。源码中darkMode和notifications使用remember保存为页面本地状态，SettingsRepository继承BaseRepository但尚未持久化业务数据。设置页目前更多承担功能入口展示和交互样式验证作用：深色模式开关、通知开关、缓存清理文字、语言切换、关于和反馈入口均有明确UI位置，但没有与数据库、SharedPreferences或系统设置建立完整数据链路。该模块的实现状态在报告中作为项目不足说明，后续可把设置状态写入DataStore或Room，并把缓存清理入口连接到CacheDao的删除语句。"),
        ("2.2.13 个人中心模块", "个人中心由ProfileScreen、ProfileSetupScreen、ProfileEditScreen和ProfileViewModel构成。注册后进入ProfileSetupScreen完善资料，已有用户可从ProfileScreen进入编辑页。ProfileViewModel订阅UserRepository.observeCurrentUser并把Room中的UserEntity映射到ProfileFormState，支持昵称、性别、生日、签名、城市和旅行偏好更新。头像支持拍照和相册选择：拍照通过ActivityResultContracts.TakePicture与FileProvider生成应用私有目录下的图片Uri，相册选择通过GetContent读取图片并复制到filesDir/profile_images目录。保存时，ProfileFormState.toInputOrNull校验昵称为2到20位中文、英文或数字，再调用UserRepository.saveUser或updateProfile写入users表。个人中心还提供收藏、历史、订单占位、设置、关于和退出登录入口。"),
        ("2.2.14 图片加载模块", "图片加载由component.Components.kt中的NetworkImage统一封装，底层使用Coil的AsyncImage。组件接收url、contentDescription、modifier和fallbackUrl，内部维护useFallback状态；当图片加载失败且允许使用兜底图片时，onError会切换到fallbackUrl。系统图片来源包括PexelsImageResponse中的城市图片、高德POI照片、高德静态地图链接以及源码中预置的Unsplash图片列表。首页的城市横幅、景点卡片、酒店和美食卡片、详情页头图、登录页背景、个人头像均复用NetworkImage。ImageRepository还会把Pexels搜索结果写入cached_images表，缓存字段包含keyword、imageUrl、photographer和cachedAt，缓存时长为24小时。"),
        ("2.2.15 网络请求模块", "网络请求分为AI请求和第三方请求两组。AIClientFactory根据ApiConfig.AI_BASE_URL创建Retrofit，OkHttpClient设置连接、写入、读取和总调用超时，并通过NetworkInterceptor添加Bearer认证头和Content-Type。OpenAIChatApi定义POST chat/completions接口，ChatRequest和ChatResponse使用kotlinx.serialization标注。ThirdPartyClientFactory复用一个OkHttpClient，根据天气、高德和Pexels的baseUrl分别创建WeatherApi、MapApi和ImageApi。Repository层统一以AppResult返回Success、Error、Loading或Empty，错误由AIErrorMapper和ThirdPartyErrorMapper转换为中文提示。UI层不直接处理Retrofit异常，只订阅ViewModel状态，网络失败时展示errorMessage。"),
        ("2.2.16 数据缓存模块", "数据缓存主要依赖Room中的cached_weather、cached_poi和cached_images三张表。WeatherRepository以城市和日期组成cacheKey，缓存有效期为30分钟；MapRepository以城市和分类组成cacheKey，缓存景点、酒店和美食POI，有效期为24小时；ImageRepository以keyword为主键缓存Pexels图片，有效期同样为24小时。CacheDao提供getWeather、upsertWeather、getPois、deletePois、upsertPois、getImage和upsertImage等方法。缓存策略没有引入复杂的同步框架，而是在Repository读取网络前完成新鲜度判断，命中缓存直接返回AppResult.Success。该设计降低了首页、地图页和天气页的重复请求次数，也使网络异常时页面具有更稳定的数据来源。"),
    ]
    for title, body in modules:
        add_title(doc, title, 3)
        add_para(doc, body)

    add_title(doc, "2.3 非功能需求", 2)
    add_para(doc, "系统的非功能需求体现在响应速度、可靠性、可维护性、可扩展性和安全性等方面。响应速度上，HomeDataViewModel和MapViewModel使用coroutineScope与async并行加载天气、POI、酒店、美食和图片，避免串行请求拉长首屏等待。可靠性上，AIErrorMapper、ThirdPartyErrorMapper和AppResult把网络异常转换为页面可展示的中文消息；WeatherRepository还在和风天气失败后回退到高德天气。可维护性上，页面、ViewModel、Repository、网络模型和数据库实体基本按功能拆分，Hilt降低构造依赖成本。可扩展性上，AppRoute集中定义路由，AIModels和ThirdPartyModels均使用数据类承载接口结构，后续新增服务时可延续现有模式。安全性上，项目已通过FileProvider处理头像文件共享，但API Key硬编码、密码明文保存和启动即请求旧相册权限仍需优化。")
    add_para(doc, "源码还体现出对用户体验连续性的要求。首页、地图、天气和行程页面都能围绕同一个城市参数继续展开，AI计划生成后可进入历史记录，城市浏览后可进入收藏，头像文件也通过应用私有目录沉淀下来。这些设计保证用户操作不是孤立页面行为，而是在本地状态、数据库和导航栈之间形成可追踪的使用链路。")

    add_title(doc, "第三章 内容与方法", 1)
    add_title(doc, "3.1 系统总体架构", 2)
    add_para(doc, "本项目采用以MVVM为核心的分层结构。UI层由Compose页面组成，包括HomeScreen、PlannerScreen、ItineraryScreen、WeatherScreen、MapScreen、FavoriteScreen、HistoryScreen、ProfileScreen和SettingsScreen等；状态层由多个HiltViewModel承载，如AIViewModel、HomeDataViewModel、MapViewModel、WeatherViewModel、ProfileViewModel和HistoryViewModel；数据层由Repository封装，包括AIRepository、GeneratedPlanRepository、WeatherRepository、MapRepository、ImageRepository、UserRepository、FavoriteRepository和HistoryRepository；底层资源包括Room数据库、SharedPreferences会话、Retrofit网络接口、Android定位服务、FileProvider和Coil图片加载。")
    add_figure_placeholder(doc, "图3-1 系统总体架构图")
    add_para(doc, "Hilt是项目连接各层的关键。AITravelMateApplication标注@HiltAndroidApp，MainActivity标注@AndroidEntryPoint，AppModules.kt中提供NetworkModule、DatabaseModule和RepositoryModule。DatabaseModule通过Room.databaseBuilder创建ai_travelmate.db并注册MIGRATION_1_2到MIGRATION_7_8；RepositoryModule向应用范围提供登录、首页、规划、收藏、历史、资料等Repository；AI和第三方Repository通过@Inject构造函数注入。项目的依赖注入不是装饰性配置，实际页面中的hiltViewModel能够直接拿到已组装好的Repository链路。")
    add_figure_placeholder(doc, "图3-2 功能模块图")
    add_figure_placeholder(doc, "图3-3 MVVM架构图")

    add_title(doc, "3.2 系统设计", 2)
    add_para(doc, "系统页面导航由AppRoute和AppNavHost统一组织。AppRoute使用sealed class定义Splash、Login、Register、CompleteProfile、Home、Planner、Itinerary、Scenic、Weather、Map、Favorite、History、Profile、EditProfile、Settings和Chat等路由，并在需要传参的页面中提供createRoute方法。AppNavHost使用NavHost注册各Composable页面，并根据当前路由的showBottomBar属性决定是否显示TravelBottomBar。底部导航包含首页、AI规划、收藏、历史和我的五个主入口，覆盖了项目的核心使用路径。")
    add_para(doc, "UI设计上，项目把按钮、卡片、顶部栏、底部栏、图片、横幅、时间线、天气卡片等组件沉淀在Components.kt中。页面不是简单重复堆叠控件，而是复用TravelBanner、CommonCard、BudgetCard、ScenicCard、TravelTimeline、NetworkImage等组件形成一致的视觉风格。theme包提供Color、Shape、Spacing和Typography，页面中通过TravelSpacing.medium、MaterialTheme.shapes.large、TravelBlue、TravelTeal等统一控制间距、圆角和色彩。该设计让不同功能模块之间保持较高一致性，也降低后续调整样式的成本。")
    add_para(doc, "项目目录职责如下：app为唯一Android模块，保存构建脚本、资源和源码；ai包负责OpenAI兼容接口、Prompt、AI模型、AIRepository和AIViewModel；thirdparty包负责和风天气、高德地图、Pexels图片、定位读取与缓存映射；database包负责Room数据库、DAO、缓存实体和迁移；model包存放Room业务实体；repository包处理用户、会话、收藏和历史；navigation包集中定义路由与NavHost；ui包按页面分组；component包提供通用Compose组件；theme包维护视觉规范；config包保存接口配置；common包定义AppResult；local和remote包目前是预留接口。")
    add_para(doc, "从模块依赖方向看，项目基本保持了由页面到状态、由状态到业务、由业务到底层资源的单向调用。ui包中的页面只通过ViewModel暴露的函数触发动作，不直接创建Retrofit、Room或OkHttp对象；ViewModel只负责组装页面状态和启动协程，不直接拼接网络路径；Repository才处理接口配置、缓存有效期、异常映射和数据库读写。这样的依赖方向使页面替换和数据源替换互不干扰。例如首页可以继续使用HomeDataViewModel的city、weather、scenic、hotels、foods和cityImage状态，而底层图片服务从Pexels迁移到其他图片源时，只需调整ImageRepository和ImageApi。")

    add_title(doc, "3.3 数据库设计", 2)
    add_para(doc, "数据库由AppDatabase.kt定义，版本号为8，exportSchema开启，因此app/schemas目录中保留了从1到8的结构快照。数据库实体包括UserEntity、City、TravelPlan、DayPlan、Scenic、Hotel、Weather、Food、Budget、ChatMessage、Favorite、History、CachedWeatherEntity、CachedPoiEntity和CachedImageEntity。实际运行中，用户、收藏、历史和缓存实体承担主要业务，城市、旅行计划、日计划、景点、酒店、天气、美食、预算和聊天消息实体为后续本地结构化扩展保留了表结构。")
    add_figure_placeholder(doc, "图3-5 数据库结构图")
    add_table(doc, ["实体", "表名", "主要作用"], [
        ["UserEntity", "users", "保存手机号、密码、昵称、头像、性别、生日、签名、城市和旅行偏好"],
        ["Favorite", "favorites", "按userPhone保存用户收藏城市，记录标题、图片和创建时间"],
        ["History", "history", "保存AI生成旅行计划摘要和planJson，用于历史恢复"],
        ["CachedWeatherEntity", "cached_weather", "缓存天气、温度、湿度、风力、空气质量和降水"],
        ["CachedPoiEntity", "cached_poi", "缓存高德POI的名称、类型、地址、经纬度、评分和图片"],
        ["CachedImageEntity", "cached_images", "缓存Pexels按关键词检索到的图片链接和摄影师"],
        ["City等基础实体", "cities等", "提供城市、路线、景点、酒店、天气、美食、预算、聊天消息等结构化表"],
    ], widths=[3, 3, 9])
    add_para(doc, "DAO层集中在Dao.kt中。UserDao提供observeCurrentUser、observeByPhone、getByPhone、getByPhoneAndPassword、upsert、updateAvatar和updateNickname；FavoriteDao提供按用户和类型观察收藏、查询单条收藏、插入替换和删除；HistoryDao提供按用户观察历史、读取指定历史、插入替换和删除；CacheDao提供天气、POI和图片的查询与写入。其他基础DAO目前只提供getAll方法。数据库迁移脚本体现了项目迭代过程：1到2新增缓存表，2到3为cached_poi补photoUrl，3到4重建users表，4到5增加password，5到6补充favorites的用户和展示字段，6到7补充history摘要字段，7到8增加planJson。")

    add_title(doc, "3.4 网络通信设计", 2)
    add_para(doc, "项目存在两套网络结构。com.example.ai.network.NetworkContracts.kt定义了通用BaseApi、AIApi、WeatherApi、HotelApi和ScenicApi，以及以DEFAULT_BASE_URL为基础的Retrofit配置，该部分在当前页面业务中没有承载真实接口。真实请求集中在ai.network和thirdparty.network。AIClientFactory面向AI_BASE_URL创建OpenAIChatApi，ThirdPartyClientFactory面向WEATHER_BASE_URL、AMAP_BASE_URL和PEXELS_BASE_URL分别创建WeatherApi、MapApi和ImageApi。两套实现均使用Retrofit 3、OkHttp 5和kotlinx.serialization converter，并配置ignoreUnknownKeys和isLenient以兼容第三方返回字段。")
    add_figure_placeholder(doc, "图3-6 网络请求流程图")
    add_para(doc, "典型数据流为：Compose页面触发事件，ViewModel在viewModelScope中启动协程，Repository检查配置和缓存，未命中缓存时调用Retrofit接口，接口返回数据类后经Mapper或扩展函数转换为UI可用模型，再通过MutableStateFlow更新UiState，页面通过collectAsStateWithLifecycle自动重组。错误流同样统一：HttpException、SocketTimeoutException、UnknownHostException、SSLException和SerializationException被ErrorMapper转换为中文说明，ViewModel写入errorMessage，页面展示错误提示。")
    add_figure_placeholder(doc, "图3-4 数据流图")

    add_title(doc, "3.5 AI功能设计", 2)
    add_para(doc, "AI能力以AIRepository为中心。ApiConfig配置AI_BASE_URL、AI_API_KEY、AI_MODEL、AI_TEMPERATURE和AI_MAX_TOKENS，AIClientFactory通过NetworkInterceptor加入Bearer认证头，OpenAIChatApi向chat/completions提交ChatRequest。源码没有接入Dify工作流，也没有出现DeepSeek专用SDK，实际协议为OpenAI兼容聊天补全接口。ChatRequest支持model、messages、temperature、max_tokens、response_format和stream字段，项目默认stream为false，结构化请求会设置ResponseFormat(type=json_object)。")
    add_figure_placeholder(doc, "图3-7 AI调用流程图")
    add_para(doc, "PromptManager是AI功能稳定性的关键。chatPrompt要求旅行助手回答简洁实用；travelPlanPrompt要求只返回合法JSON对象，并逐项声明TravelPlan所需字段；weatherPrompt要求返回outfit、umbrella、sunscreen、outdoor；budgetPrompt要求返回reasonable、summary、hotelAdvice、trafficAdvice、foodAdvice、savingAdvice；guideSummaryPrompt要求返回scenicSpots、routes、foods、warnings。AIRepository的decodeTravelPlan先直接反序列化，失败后调用normalizeTravelPlanJson统一字段名和默认值，解决AI返回字段不完全一致的问题。AIViewModel把聊天、重生成、删除消息、清空会话、生成旅行计划、天气建议、预算分析和攻略总结都收束到同一个AIUiState中。")

    add_title(doc, "3.6 核心代码分析", 2)
    add_para(doc, "AIRepository体现了本项目对大模型结果的工程化处理。generateTravelPlan并不把AI返回文本直接传给页面，而是用sendRaw发送消息、要求JSON模式、限制TRAVEL_PLAN_MAX_TOKENS为2400，再用decodeTravelPlan得到TravelPlan对象。normalizeDays、normalizeSchedules、normalizeBudget、normalizeHotels、normalizeFoods和normalizeTips把常见字段变体转为项目定义字段，numberValue还会从文本中提取数字。这种处理降低了大模型输出轻微偏差对UI的影响，使ItineraryScreen可以稳定读取days、schedules、budget、hotels、foods和tips。")
    add_para(doc, "ThirdPartyRepositories.kt体现了第三方服务的封装边界。WeatherRepository先查缓存，再优先访问和风天气，失败后调用高德天气；MapRepository把景点、美食、酒店统一抽象为searchPoi，并把高德AMapPoi映射为内部Poi；ImageRepository负责Pexels图片检索和缓存；LocationRepository负责Android定位权限检查、读取最近位置或单次更新，并调用高德逆地理编码得到城市名。页面层完全不关心接口路径、Key、缓存有效期和异常类型，只拿到AppResult后的领域模型。")
    add_para(doc, "GeneratedPlanRepository连接了AI生成和历史记录。其latestPlan是StateFlow，ItineraryScreen订阅后可直接展示最新AI方案；updatePlan在更新latestPlan的同时把完整TravelPlan序列化为planJson，并记录天数、行程数、酒店数和美食数。HistoryScreen再次打开历史时，openHistoryPlan从Room读取History，再反序列化planJson回TravelPlan。该实现避免了只保存摘要导致历史无法复原的问题，也让历史页成为AI规划结果的长期入口。")
    add_para(doc, "用户体系由SessionRepository、UserRepository和UserDao支撑。SessionRepository使用SharedPreferences保存当前手机号和记住登录标志，并通过MutableStateFlow暴露currentPhone。UserRepository基于currentPhone切换observeCurrentUser的数据源，未登录时返回空Flow，已登录时观察UserDao.observeByPhone。收藏和历史也采用同样的用户隔离方式：FavoriteRepository与HistoryRepository读取当前手机号，只返回当前用户的数据。该设计在没有远程账号系统的前提下，完成了设备端多用户数据隔离。")

    add_title(doc, "3.7 系统测试", 2)
    add_para(doc, "本次报告生成前对项目执行了Gradle单元测试任务：.\\gradlew.bat testDebugUnitTest。构建过程完成了Debug资源处理、Kotlin编译、Kapt处理、Hilt聚合、Room注解处理和debug单元测试执行，结果为BUILD SUCCESSFUL，用时约49秒。输出中存在Android Gradle Plugin关于android.builtInKotlin和android.newDsl的弃用警告，以及Kapt阶段若干处理器未识别选项警告，但没有造成编译或测试失败。项目当前包含ExampleUnitTest和ExampleInstrumentedTest模板，测试覆盖仍以构建连通性为主，尚未对Repository缓存策略、AI JSON归一化、用户登录和历史恢复建立专门单元测试。")
    add_para(doc, "从功能测试角度，系统应重点验证以下流程：首次启动进入Splash后根据记住登录状态跳转；注册验证码、本地资料保存和登录验证；首页定位失败后的城市搜索；城市加载后天气、景点、酒店、美食和图片并行展示；AI规划成功后跳转行程页并保存历史；历史记录可重新打开完整计划；收藏城市后FavoriteScreen实时刷新；天气页完成真实天气和AI建议展示；地图页能够打开高德地图或Web链接。源码中的状态流和Repository边界使这些测试均具备可自动化的入口。")

    add_title(doc, "3.8 项目运行效果", 2)
    add_para(doc, "系统运行后，用户首先看到登录或注册页面，完成账号流程后进入首页。首页以用户信息、搜索城市、城市横幅、收藏动作、快捷AI规划、天气摘要、景点酒店美食列表组成首屏和滚动内容。AI规划页提供日期选择器、预算输入、预算分项、出行人数和兴趣标签，生成后进入行程页。行程页按天展示时间线，并显示预算、住宿和活动数量。天气页展示实时天气、空气质量、湿度、风力、降雨和AI建议。地图页按景点、酒店、美食分组展示POI，点击可打开高德地图路线。收藏页与历史页分别展示用户主动保存的城市和AI生成过的计划。")
    add_figure_placeholder(doc, "图3-8 系统运行界面")

    add_title(doc, "3.8.1 项目亮点", 3)
    highlights = [
        "采用Compose声明式UI，所有主要页面均由Composable函数实现，组件复用程度较高。",
        "MVVM与Repository分层清晰，页面事件、状态维护、业务封装和底层资源访问职责明确。",
        "Hilt贯穿Application、Activity、ViewModel、Repository和Database，降低对象装配复杂度。",
        "AI规划不是文本展示，而是通过Prompt约束和JSON解析形成可展示、可保存、可恢复的TravelPlan。",
        "天气、地图、图片三类第三方服务均有独立Repository，并且包含缓存和错误中文化处理。",
        "Room数据库不仅保存用户资料，还保存收藏、历史和缓存，支撑离线数据沉淀。",
        "首页加载采用协程并行请求，天气、POI、酒店、美食和图片能够同时拉取。",
        "历史记录保存完整planJson，用户可从历史页重新打开AI生成的行程。",
        "NetworkImage封装Coil加载和兜底图片，避免单个图片失败破坏页面展示。",
        "地图模块兼容高德App深链和Web链接，未内嵌地图SDK也能完成路线入口。",
    ]
    for item in highlights:
        add_para(doc, item, first_line=False)

    add_title(doc, "3.8.2 项目不足与优化方案", 3)
    weaknesses = [
        "API Key写在ApiConfig.kt源码常量中，存在泄露风险。优化方案是将密钥迁移到local.properties、CI环境变量或后端代理服务，客户端只保存受限配置。",
        "用户密码以明文形式保存在users表中。优化方案是引入加盐哈希或接入后端认证服务，本地只保存会话令牌。",
        "注册验证码由前端随机生成并显示在页面中，只适合开发演示。优化方案是接入短信网关或服务端验证码接口，并设置验证码过期和频率限制。",
        "MainActivity启动时请求READ_EXTERNAL_STORAGE，且该权限在高版本Android中已被新媒体权限替代。优化方案是移除启动即请求，改为在选择头像时按系统版本请求最小权限。",
        "SettingsScreen中的深色模式、通知和缓存清理尚未持久化。优化方案是使用DataStore保存设置，并把清理按钮连接到CacheDao删除接口。",
        "Room表之间没有外键和索引约束，favorites和history主要依赖字符串字段关联用户。优化方案是在用户手机号、targetType、targetId、visitedAt等字段上增加索引，并明确级联删除规则。",
        "地图功能依赖外部高德App或Web页，没有应用内地图视图。优化方案是接入高德地图Android SDK，展示标记、路线规划和多点行程路径。",
        "AI请求为非流式返回，长行程生成时用户只能等待完整结果。优化方案是启用stream，按片段更新UI，并在生成完成后再执行JSON结构化校验。",
        "测试覆盖仍停留在模板和构建连通性。优化方案是为AIRepository的normalize逻辑、缓存过期判断、UserRepository登录和GeneratedPlanRepository历史恢复编写单元测试。",
        "local和remote包、通用NetworkContracts以及部分基础Repository仍处于预留状态。优化方案是清理未使用接口，或把通用网络层与实际第三方网络层合并成一致结构。",
    ]
    for item in weaknesses:
        add_para(doc, item, first_line=False)

    add_title(doc, "第四章 总结", 1)
    add_para(doc, "WeTravel项目完成了一个较完整的Android智能旅行规划系统。源码显示，项目以Compose和Navigation构建多页面体验，以ViewModel和StateFlow维护响应式状态，以Repository隔离本地数据、网络请求和AI能力，以Room保存用户、收藏、历史和缓存，以Retrofit与OkHttp连接大模型、和风天气、高德地图和Pexels图片服务。项目功能覆盖登录注册、个人资料、首页聚合、城市搜索、AI规划、行程展示、天气建议、地图POI、景点酒店美食、收藏、历史和设置。")
    add_para(doc, "从课程设计角度看，本项目的价值在于把多个Android知识点放入真实业务链路中验证：Hilt不是孤立示例，而是支撑Repository和ViewModel装配；Room不是简单增删改查，而是同时承担账号、历史和缓存；Retrofit不是单一接口请求，而是面对AI、天气、地图和图片四类服务；AI不是聊天玩具，而是通过Prompt和结构化解析进入行程模型。项目仍存在密钥管理、密码安全、权限适配、设置持久化、测试覆盖和地图内嵌能力等不足，但整体架构已经具备继续迭代为完整旅行应用的基础。")

    add_title(doc, "参考文献", 1)
    refs = [
        "[1] Google. Android Developers: Jetpack Compose[EB/OL]. (2026-07-06)[2026-07-06]. https://developer.android.com/compose.",
        "[2] Google. Guide to app architecture[EB/OL]. (2026-07-06)[2026-07-06]. https://developer.android.com/topic/architecture.",
        "[3] Google. Room persistence library[EB/OL]. (2026-07-06)[2026-07-06]. https://developer.android.com/training/data-storage/room.",
        "[4] Google. Hilt dependency injection[EB/OL]. (2026-07-06)[2026-07-06]. https://developer.android.com/training/dependency-injection/hilt-android.",
        "[5] Google. Navigation with Compose[EB/OL]. (2026-07-06)[2026-07-06]. https://developer.android.com/develop/ui/compose/navigation.",
        "[6] Square. Retrofit documentation[EB/OL]. (2026-07-06)[2026-07-06]. https://square.github.io/retrofit/.",
        "[7] Square. OkHttp documentation[EB/OL]. (2026-07-06)[2026-07-06]. https://square.github.io/okhttp/.",
        "[8] Kotlin. kotlinx.serialization guide[EB/OL]. (2026-07-06)[2026-07-06]. https://github.com/Kotlin/kotlinx.serialization.",
        "[9] Coil. Coil Compose image loading[EB/OL]. (2026-07-06)[2026-07-06]. https://coil-kt.github.io/coil/compose/.",
        "[10] OpenAI. Chat Completions API reference[EB/OL]. (2026-07-06)[2026-07-06]. https://platform.openai.com/docs/api-reference/chat.",
        "[11] QWeather. Weather API documentation[EB/OL]. (2026-07-06)[2026-07-06]. https://dev.qweather.com/docs/.",
        "[12] 高德开放平台. Web服务API文档[EB/OL]. (2026-07-06)[2026-07-06]. https://lbs.amap.com/api/webservice/summary.",
    ]
    for ref in refs:
        add_para(doc, ref, first_line=False)

    add_title(doc, "学习心得", 1)
    add_para(doc, "通过对WeTravel项目的设计与实现分析，可以更加清楚地理解Android课程项目与课堂示例之间的差别。单独学习Compose、Room、Retrofit或ViewModel时，容易把技术看成彼此分离的知识点；在本项目中，用户点击一次生成旅行方案，背后会同时涉及页面输入校验、ViewModel状态更新、Prompt构造、Retrofit请求、JSON解析、Room历史写入和Navigation跳转。这个过程说明移动端开发的难点并不只是写出某个页面，而是让数据在多个模块之间稳定流动。")
    add_para(doc, "项目中让我印象最深的是AI结果结构化处理。AIRepository没有把大模型返回当作普通文本，而是通过PromptManager限制字段，再用normalizeTravelPlanJson修正字段差异，这体现了把AI能力工程化接入应用的思路。另一方面，源码也暴露出课程项目常见问题，如密钥硬编码、密码明文保存、测试不足、设置功能未持久化等。分析这些不足比单纯罗列功能更有价值，因为它们直接指向真实软件迭代中必须解决的安全性、可靠性和可维护性问题。后续继续完善该项目时，应优先加强账号安全、配置安全、自动化测试和地图体验，使系统从可运行的课程项目进一步接近可长期维护的移动应用。")

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_document()
    print(path)
