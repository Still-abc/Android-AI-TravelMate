from __future__ import annotations

import math
import re
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(r"E:\Android_work\Wetravel")
ASSET_DIR = ROOT / "report_assets"
OUT_DIR = Path.home() / "Downloads"
INPUT_NEEDLE = "\u7b2c\u4e00\u8f6e\u4fee\u6539"
OUTPUT_SUFFIX = "_第二轮修改_补图排版优化.docx"

FONT_REGULAR = r"C:\Windows\Fonts\msyh.ttc"
FONT_BOLD = r"C:\Windows\Fonts\msyhbd.ttc"
FONT_SONG = r"C:\Windows\Fonts\simsun.ttc"

W, H = 1600, 950
INK = "#17324D"
MUTED = "#607083"
BLUE = "#2E74B5"
TEAL = "#19A7A8"
CORAL = "#F26D5B"
AMBER = "#F2B84B"
GREEN = "#51A96B"
PURPLE = "#7B61B8"
BG = "#F6F8FB"
LINE = "#CED8E4"
CARD = "#FFFFFF"


def font(size: int, bold: bool = False):
    path = FONT_BOLD if bold else FONT_REGULAR
    return ImageFont.truetype(path, size=size)


def text_size(draw: ImageDraw.ImageDraw, text: str, fnt) -> tuple[int, int]:
    bbox = draw.textbbox((0, 0), text, font=fnt)
    return bbox[2] - bbox[0], bbox[3] - bbox[1]


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for ch in text:
        trial = current + ch
        if text_size(draw, trial, fnt)[0] <= max_width or not current:
            current = trial
        else:
            lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def draw_wrapped(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, fnt, fill: str, max_width: int, line_gap: int = 8, align: str = "left"):
    x, y = xy
    lines = wrap_text(draw, text, fnt, max_width)
    for line in lines:
        w, h = text_size(draw, line, fnt)
        lx = x if align == "left" else x + (max_width - w) // 2
        draw.text((lx, y), line, font=fnt, fill=fill)
        y += h + line_gap
    return y


def rounded(draw, box, fill=CARD, outline=LINE, radius=22, width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def pill(draw, box, text, fill, txt="#FFFFFF", size=28):
    draw.rounded_rectangle(box, radius=(box[3] - box[1]) // 2, fill=fill)
    f = font(size, True)
    tw, th = text_size(draw, text, f)
    draw.text(((box[0]+box[2]-tw)//2, (box[1]+box[3]-th)//2-2), text, font=f, fill=txt)


def arrow(draw, start, end, color=BLUE, width=4):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    ang = math.atan2(y2-y1, x2-x1)
    size = 16
    pts = [
        (x2, y2),
        (x2 - size * math.cos(ang - math.pi/6), y2 - size * math.sin(ang - math.pi/6)),
        (x2 - size * math.cos(ang + math.pi/6), y2 - size * math.sin(ang + math.pi/6)),
    ]
    draw.polygon(pts, fill=color)


def header(draw, title, subtitle):
    draw.text((70, 42), title, font=font(44, True), fill=INK)
    draw.text((72, 98), subtitle, font=font(24), fill=MUTED)
    draw.line((70, 140, W-70, 140), fill=LINE, width=2)


def card_with_title(draw, box, title, body: Sequence[str] | str, color=BLUE, title_size=29, body_size=22):
    rounded(draw, box, fill=CARD, outline=LINE, radius=24, width=2)
    x1, y1, x2, y2 = box
    draw.rounded_rectangle((x1, y1, x2, y1 + 58), radius=24, fill=color)
    draw.rectangle((x1, y1 + 32, x2, y1 + 58), fill=color)
    draw.text((x1 + 24, y1 + 13), title, font=font(title_size, True), fill="#FFFFFF")
    lines = [body] if isinstance(body, str) else list(body)
    y = y1 + 82
    for item in lines:
        y = draw_wrapped(draw, (x1 + 24, y), item, font(body_size), INK, x2 - x1 - 48, 6)
        y += 5


def new_canvas() -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (W, H), BG)
    draw = ImageDraw.Draw(img)
    return img, draw


def save(img: Image.Image, name: str) -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    path = ASSET_DIR / name
    img.save(path, quality=95)
    return path


def fig_architecture() -> Path:
    img, d = new_canvas()
    header(d, "WeTravel 系统总体架构", "Compose 页面、ViewModel、Repository、Room 与第三方服务的分层关系")
    layers = [
        ("UI层", "HomeScreen / PlannerScreen / ItineraryScreen / WeatherScreen / MapScreen / ProfileScreen", BLUE),
        ("状态层", "AIViewModel / HomeDataViewModel / WeatherViewModel / MapViewModel / ProfileViewModel", TEAL),
        ("业务层", "AIRepository / WeatherRepository / MapRepository / ImageRepository / UserRepository", PURPLE),
        ("数据与服务层", "Room(AppDatabase) / SharedPreferences / Retrofit+OkHttp / LocationManager / FileProvider", CORAL),
    ]
    y = 175
    boxes = []
    for title, body, color in layers:
        box = (110, y, 1490, y + 118)
        card_with_title(d, box, title, body, color=color, title_size=30, body_size=25)
        boxes.append(box)
        y += 165
    for a, b in zip(boxes, boxes[1:]):
        arrow(d, ((a[0]+a[2])//2, a[3]+10), ((b[0]+b[2])//2, b[1]-10), BLUE, 5)
    ext = [
        ("AI接口", "OpenAI兼容 chat/completions", BLUE),
        ("天气", "和风天气 + 高德天气回退", TEAL),
        ("地图", "高德POI / 地理编码 / 深链", GREEN),
        ("图片", "Pexels + Unsplash兜底", AMBER),
    ]
    x = 110
    for title, body, color in ext:
        card_with_title(d, (x, 805, x+320, 900), title, body, color=color, title_size=25, body_size=20)
        x += 350
    return save(img, "fig3-1_system_architecture.png")


def fig_modules() -> Path:
    img, d = new_canvas()
    header(d, "WeTravel 功能模块", "从账号、城市数据、AI规划到个人数据沉淀的功能闭环")
    center = (620, 360, 980, 560)
    rounded(d, center, fill="#EAF3FF", outline=BLUE, radius=38, width=4)
    d.text((697, 407), "WeTravel", font=font(48, True), fill=BLUE)
    d.text((682, 472), "AI智能旅行规划系统", font=font(25, True), fill=INK)
    modules = [
        ("登录注册", "本地账号 / 记住登录", 130, 190, BLUE),
        ("首页聚合", "天气、POI、图片、攻略", 470, 170, TEAL),
        ("AI规划", "Prompt + JSON行程", 820, 170, PURPLE),
        ("城市搜索", "城市参数驱动多服务", 1170, 190, GREEN),
        ("天气建议", "和风天气 + AI穿搭", 120, 620, TEAL),
        ("地图POI", "景点/酒店/美食", 450, 655, GREEN),
        ("收藏历史", "Room持久化", 790, 655, CORAL),
        ("个人中心设置", "资料、头像、偏好", 1130, 620, AMBER),
    ]
    for title, body, x, y, color in modules:
        box = (x, y, x+300, y+130)
        card_with_title(d, box, title, body, color=color, title_size=26, body_size=21)
        arrow(d, ((box[0]+box[2])//2, (box[1]+box[3])//2), ((center[0]+center[2])//2, (center[1]+center[3])//2), color, 3)
    return save(img, "fig3-2_function_modules.png")


def fig_mvvm() -> Path:
    img, d = new_canvas()
    header(d, "MVVM 架构关系", "页面事件向下传递，StateFlow状态向上刷新")
    columns = [
        ("View", ["Compose页面", "TravelBanner / CommonCard", "Navigation Compose"], BLUE),
        ("ViewModel", ["MutableStateFlow", "viewModelScope协程", "UI状态组装"], TEAL),
        ("Repository", ["业务规则封装", "缓存判断", "错误中文化"], PURPLE),
        ("Model/Data", ["Room实体与DAO", "Retrofit数据类", "第三方API响应"], CORAL),
    ]
    x = 85
    boxes = []
    for title, items, color in columns:
        box = (x, 235, x+330, 675)
        rounded(d, box, fill=CARD, outline=color, radius=28, width=4)
        pill(d, (x+45, 270, x+285, 325), title, color, size=28)
        yy = 375
        for item in items:
            rounded(d, (x+35, yy, x+295, yy+62), fill="#F7FAFC", outline=LINE, radius=18, width=2)
            d.text((x+58, yy+17), item, font=font(22), fill=INK)
            yy += 92
        boxes.append(box)
        x += 380
    for a, b in zip(boxes, boxes[1:]):
        arrow(d, (a[2]+20, 405), (b[0]-20, 405), BLUE, 5)
        arrow(d, (b[0]-20, 530), (a[2]+20, 530), TEAL, 5)
    d.text((480, 760), "事件/请求", font=font(26, True), fill=BLUE)
    d.text((880, 760), "AppResult / StateFlow / UI刷新", font=font(26, True), fill=TEAL)
    return save(img, "fig3-3_mvvm.png")


def fig_data_flow() -> Path:
    img, d = new_canvas()
    header(d, "系统数据流", "用户输入、缓存、网络和页面状态之间的闭环")
    steps = [
        ("用户操作", "搜索城市 / 生成行程 / 收藏 / 打开历史", BLUE),
        ("ViewModel", "校验输入并启动协程", TEAL),
        ("Repository", "读取会话、判断缓存、组织请求", PURPLE),
        ("数据源", "Room缓存或Retrofit网络", CORAL),
        ("状态刷新", "StateFlow更新UiState", GREEN),
        ("Compose重组", "首页、天气、地图、行程页展示", AMBER),
    ]
    y = 205
    prev = None
    for idx, (title, body, color) in enumerate(steps):
        x = 145 if idx % 2 == 0 else 770
        box = (x, y, x+545, y+110)
        card_with_title(d, box, title, body, color=color, title_size=25, body_size=22)
        if prev:
            arrow(d, (prev[2], (prev[1]+prev[3])//2), (box[0], (box[1]+box[3])//2), BLUE if idx % 2 else TEAL, 4)
        prev = box
        y += 125
    arrow(d, (1030, 815), (360, 815), GREEN, 5)
    d.text((516, 838), "历史记录与收藏继续沉淀到Room", font=font(25, True), fill=GREEN)
    return save(img, "fig3-4_data_flow.png")


def fig_database() -> Path:
    img, d = new_canvas()
    header(d, "Room 数据库结构", "AppDatabase version=8，用户、收藏、历史与缓存为运行主表")
    tables = [
        ("users", ["id PK", "phone", "password", "nickname", "avatar", "travelPreference"], BLUE, 95, 190),
        ("favorites", ["id PK", "userPhone", "targetId", "targetType", "title", "imageUrl"], TEAL, 490, 190),
        ("history", ["id PK", "userPhone", "targetId", "days", "scheduleCount", "planJson"], PURPLE, 885, 190),
        ("cached_weather", ["cacheKey PK", "city", "date", "weather", "temperature", "cachedAt"], CORAL, 1280, 190),
        ("cached_poi", ["id PK", "cacheKey", "category", "name", "lon/lat", "photoUrl"], GREEN, 300, 610),
        ("cached_images", ["keyword PK", "imageUrl", "photographer", "cachedAt"], AMBER, 700, 610),
        ("预留实体表", ["cities", "travel_plans", "day_plans", "scenic_spots", "hotels", "foods"], "#8393A7", 1100, 610),
    ]
    for name, fields, color, x, y in tables:
        box = (x, y, x+270, y+285)
        rounded(d, box, fill=CARD, outline=color, radius=18, width=3)
        d.rounded_rectangle((x, y, x+270, y+54), radius=18, fill=color)
        d.rectangle((x, y+30, x+270, y+54), fill=color)
        d.text((x+18, y+13), name, font=font(23, True), fill="#FFFFFF")
        yy = y+72
        for f in fields:
            d.text((x+24, yy), f, font=font(20), fill=INK)
            yy += 32
    arrow(d, (365, 475), (535, 335), BLUE, 3)
    arrow(d, (365, 475), (930, 335), BLUE, 3)
    d.text((610, 500), "userPhone实现用户数据隔离", font=font(24, True), fill=BLUE)
    return save(img, "fig3-5_database.png")


def fig_network() -> Path:
    img, d = new_canvas()
    header(d, "网络请求流程", "Retrofit + OkHttp + kotlinx.serialization + AppResult")
    left = [
        ("Compose页面", "HomeScreen / WeatherScreen / MapScreen / PlannerScreen", BLUE),
        ("ViewModel", "viewModelScope.launch / async并行请求", TEAL),
        ("Repository", "缓存优先、配置检查、异常映射", PURPLE),
        ("Retrofit接口", "OpenAIChatApi / WeatherApi / MapApi / ImageApi", CORAL),
    ]
    y = 190
    boxes = []
    for title, body, color in left:
        box = (115, y, 675, y+105)
        card_with_title(d, box, title, body, color=color, title_size=24, body_size=20)
        boxes.append(box)
        y += 150
    for a, b in zip(boxes, boxes[1:]):
        arrow(d, ((a[0]+a[2])//2, a[3]+8), ((b[0]+b[2])//2, b[1]-8), BLUE, 4)
    apis = [
        ("AI服务", "chat/completions", BLUE),
        ("和风天气", "lookup/now/forecast/air", TEAL),
        ("高德地图", "POI/地理编码/天气/静态图", GREEN),
        ("Pexels", "v1/search 图片检索", AMBER),
    ]
    y = 200
    for title, body, color in apis:
        card_with_title(d, (925, y, 1465, y+105), title, body, color=color, title_size=24, body_size=20)
        arrow(d, (690, y+52), (910, y+52), color, 4)
        y += 150
    card_with_title(d, (925, 800, 1465, 890), "统一返回", "Success / Error / Loading / Empty -> UiState", color="#8393A7", title_size=24, body_size=20)
    arrow(d, (1195, 760), (1195, 795), "#8393A7", 4)
    return save(img, "fig3-6_network.png")


def fig_ai_flow() -> Path:
    img, d = new_canvas()
    header(d, "AI 调用流程", "PromptManager约束JSON，AIRepository解析并归一化为领域模型")
    steps = [
        ("PlannerScreen", "目的地、日期、预算、人数、兴趣", BLUE),
        ("AIViewModel", "generateTravelPlan / sendMessage", TEAL),
        ("PromptManager", "system + user，强制合法JSON字段", PURPLE),
        ("AIRepository", "ChatRequest + json_object", CORAL),
        ("OpenAI兼容接口", "Authorization: Bearer，chat/completions", GREEN),
        ("结果处理", "decodeTravelPlan + normalizeTravelPlanJson", AMBER),
        ("页面与历史", "ItineraryScreen展示，History保存planJson", BLUE),
    ]
    x, y = 115, 185
    boxes = []
    for i, (title, body, color) in enumerate(steps):
        box = (x, y, x+610, y+92) if i % 2 == 0 else (875, y, 1485, y+92)
        card_with_title(d, box, title, body, color=color, title_size=23, body_size=19)
        boxes.append(box)
        if i % 2 == 1:
            y += 130
    for a, b in zip(boxes, boxes[1:]):
        arrow(d, (a[2] if a[0] < b[0] else a[0], (a[1]+a[3])//2), (b[0] if a[0] < b[0] else b[2], (b[1]+b[3])//2), BLUE, 4)
    rounded(d, (100, 775, 1500, 880), fill="#EAF3FF", outline=BLUE, radius=24, width=2)
    d.text((135, 805), "结构化字段：title / departure / destination / days / schedules / budget / hotels / foods / tips", font=font(27, True), fill=BLUE)
    return save(img, "fig3-7_ai_flow.png")


def phone(draw, box, title, accent, lines):
    x1, y1, x2, y2 = box
    rounded(draw, box, fill="#111827", outline="#0B1220", radius=34, width=3)
    screen = (x1+16, y1+28, x2-16, y2-28)
    rounded(draw, screen, fill="#F8FAFC", outline="#E5E7EB", radius=24, width=2)
    draw.rounded_rectangle((x1+78, y1+13, x2-78, y1+23), radius=6, fill="#374151")
    draw.text((screen[0]+22, screen[1]+22), title, font=font(25, True), fill=accent)
    yy = screen[1]+70
    for label, value, color in lines:
        rounded(draw, (screen[0]+20, yy, screen[2]-20, yy+72), fill="#FFFFFF", outline="#E2E8F0", radius=16, width=2)
        draw.text((screen[0]+38, yy+14), label, font=font(18, True), fill=color)
        draw.text((screen[0]+38, yy+42), value, font=font(15), fill=MUTED)
        yy += 88


def fig_runtime_ui() -> Path:
    img, d = new_canvas()
    header(d, "系统运行界面示意", "根据源码页面结构绘制：首页、AI规划、行程、天气与地图核心界面")
    phones = [
        ("首页", TEAL, [("城市横幅", "天气 + POI + 图片", TEAL), ("热门景点", "高德POI横向卡片", BLUE), ("快捷AI规划", "生成推荐攻略", PURPLE)]),
        ("AI规划", PURPLE, [("基础信息", "目的地/日期", BLUE), ("预算范围", "总预算与分项", CORAL), ("旅行兴趣", "美食/自然/摄影", TEAL)]),
        ("行程", BLUE, [("行程概览", "预算/住宿/活动", BLUE), ("Day 1", "时间线Schedule", TEAL), ("旅行提示", "酒店/美食/Tips", AMBER)]),
        ("天气", TEAL, [("实时天气", "温度/湿度/风力", TEAL), ("AI提醒", "穿搭/雨伞/防晒", PURPLE), ("详情卡片", "空气质量/降雨", BLUE)]),
        ("地图", GREEN, [("城市POI", "景点/酒店/美食", GREEN), ("地点卡片", "地址/评分/电话", BLUE), ("打开路线", "高德App或Web", CORAL)]),
    ]
    x = 80
    for title, accent, lines in phones:
        phone(d, (x, 190, x+270, 860), title, accent, lines)
        x += 300
    return save(img, "fig3-8_runtime_ui.png")


FIGURES = {
    "\u56fe3-1 \u7cfb\u7edf\u603b\u4f53\u67b6\u6784\u56fe": fig_architecture,
    "\u56fe3-2 \u529f\u80fd\u6a21\u5757\u56fe": fig_modules,
    "\u56fe3-3 MVVM\u67b6\u6784\u56fe": fig_mvvm,
    "\u56fe3-5 \u6570\u636e\u5e93\u7ed3\u6784\u56fe": fig_database,
    "\u56fe3-6 \u7f51\u7edc\u8bf7\u6c42\u6d41\u7a0b\u56fe": fig_network,
    "\u56fe3-4 \u6570\u636e\u6d41\u56fe": fig_data_flow,
    "\u56fe3-7 AI\u8c03\u7528\u6d41\u7a0b\u56fe": fig_ai_flow,
    "\u56fe3-8 \u7cfb\u7edf\u8fd0\u884c\u754c\u9762": fig_runtime_ui,
}


def set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=None, bold=None, color=None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_format(paragraph, style_name: str):
    pf = paragraph.paragraph_format
    if style_name == "Normal":
        text = paragraph.text.strip()
        pf.line_spacing = 1.25
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        if len(text) > 40 and not re.match(r"^\[\d+\]", text):
            pf.first_line_indent = Cm(0.74)
        for run in paragraph.runs:
            set_run_font(run, east_asia="宋体", size=11, color="222222")
    elif style_name == "Heading 1":
        pf.space_before = Pt(14)
        pf.space_after = Pt(8)
        pf.keep_with_next = True
        for run in paragraph.runs:
            set_run_font(run, east_asia="黑体", ascii_font="Arial", size=16, bold=True, color="1F4D78")
    elif style_name == "Heading 2":
        pf.space_before = Pt(10)
        pf.space_after = Pt(6)
        pf.keep_with_next = True
        for run in paragraph.runs:
            set_run_font(run, east_asia="黑体", ascii_font="Arial", size=13.5, bold=True, color="2E74B5")
    elif style_name == "Heading 3":
        pf.space_before = Pt(8)
        pf.space_after = Pt(4)
        pf.keep_with_next = True
        for run in paragraph.runs:
            set_run_font(run, east_asia="黑体", ascii_font="Arial", size=12, bold=True, color="1F4D78")


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph(text)
    new_para._p = new_p
    new_para._element = new_p
    if style is not None:
        new_para.style = style
    return new_para


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def format_tables(doc: Document):
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for r_idx, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.space_after = Pt(2)
                    for run in p.runs:
                        set_run_font(run, east_asia="宋体", size=10.5, color="222222")
                if r_idx == 0:
                    set_cell_shading(cell, "E8EEF5")
                    for p in cell.paragraphs:
                        for run in p.runs:
                            set_run_font(run, east_asia="黑体", ascii_font="Arial", size=10.5, bold=True, color="1F4D78")


def find_input() -> Path:
    files = [p for p in OUT_DIR.glob("*.docx") if "WeTravel" in p.name and INPUT_NEEDLE in p.name and not p.name.startswith("~$")]
    if not files:
        raise FileNotFoundError("Cannot locate first-round WeTravel docx in Downloads")
    return max(files, key=lambda p: p.stat().st_mtime)


def apply_styles(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.35)
    section.bottom_margin = Cm(2.35)
    section.left_margin = Cm(2.65)
    section.right_margin = Cm(2.65)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)

    for p in doc.paragraphs:
        if p.style.name in {"Normal", "Heading 1", "Heading 2", "Heading 3"}:
            set_paragraph_format(p, p.style.name)

    if len(doc.paragraphs) >= 2:
        for idx, size in [(0, 22), (1, 18)]:
            p = doc.paragraphs[idx]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(24 if idx == 0 else 34)
            for run in p.runs:
                set_run_font(run, east_asia="黑体", ascii_font="Arial", size=size, bold=True, color="1F4D78")

    format_tables(doc)


def replace_placeholders(doc: Document, images: dict[str, Path]):
    for p in list(doc.paragraphs):
        raw = p.text.strip()
        caption = raw.strip("【】")
        if caption in images:
            p.clear()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run()
            run.add_picture(str(images[caption]), width=Cm(15.6))
            p._p.get_or_add_pPr().append(OxmlElement("w:keepNext"))
            cap = insert_paragraph_after(p, caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.first_line_indent = Cm(0)
            cap.paragraph_format.space_before = Pt(2)
            cap.paragraph_format.space_after = Pt(10)
            for run in cap.runs:
                set_run_font(run, east_asia="宋体", size=10.5, bold=True, color="4B5563")


def add_footer(doc: Document):
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "基于Android的AI智能旅行规划系统（WeTravel）的设计与实现"
    for run in footer.runs:
        set_run_font(run, east_asia="宋体", size=9, color="6B7280")


def generate_assets() -> dict[str, Path]:
    return {caption: make() for caption, make in FIGURES.items()}


def main():
    input_path = find_input()
    images = generate_assets()
    doc = Document(input_path)
    apply_styles(doc)
    replace_placeholders(doc, images)
    add_footer(doc)
    output = input_path.with_name(input_path.stem + OUTPUT_SUFFIX)
    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()
